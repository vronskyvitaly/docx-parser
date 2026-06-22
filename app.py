"""
docx-parser — HTTP-сервис для:
  • извлечения текста из DOCX  (POST /extract)
  • генерации договора Royal Cargo в PDF (POST /generate-contract)
"""
import io
import os
import logging
import subprocess
import tempfile
from pathlib import Path

from flask import Flask, request, jsonify, send_file
from docx import Document
from docxtpl import DocxTemplate

logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s %(message)s')
log = logging.getLogger(__name__)

app = Flask(__name__)

# === Конфиг ===
TEMPLATE_PATH = os.environ.get(
    'CONTRACT_TEMPLATE_PATH',
    '/app/templates/contract.docx'
)

# Реквизиты Royal Cargo. Можно переопределить через env vars в Coolify.
# Дефолты — заглушки, ОБЯЗАТЕЛЬНО заменить на реальные перед прод-использованием.
COMPANY = {
    'name':              os.environ.get('ROYAL_NAME',          'ООО «Роял Карго»'),
    'inn':               os.environ.get('ROYAL_INN',           '0000000000'),
    'kpp':               os.environ.get('ROYAL_KPP',           '000000000'),
    'ogrn':              os.environ.get('ROYAL_OGRN',          '0000000000000'),
    'address':           os.environ.get('ROYAL_ADDRESS',       'г. Москва'),
    'account':           os.environ.get('ROYAL_ACCOUNT',       '00000000000000000000'),
    'bank':              os.environ.get('ROYAL_BANK',          'Банк'),
    'bik':               os.environ.get('ROYAL_BIK',           '000000000'),
    'cor_account':       os.environ.get('ROYAL_COR_ACCOUNT',   '00000000000000000000'),
    'director_genitive': os.environ.get('ROYAL_DIRECTOR_GEN',  'Директора Директора Директоровича'),
    'director_short':    os.environ.get('ROYAL_DIRECTOR_SHORT','Директоров Д.Д.'),
    'email':             os.environ.get('ROYAL_EMAIL',         'cargo@tamozhennyy.broker'),
    'phone':             os.environ.get('ROYAL_PHONE',         '+7 (495) 128-25-35'),
}


# ============================================================
# Health
# ============================================================

@app.get('/')
def root():
    return jsonify({"status": "ok", "service": "docx-parser",
                    "endpoints": ["/health", "/extract",
                                  "/generate-contract", "/generate-contract-docx"]})


@app.get('/health')
def health_check():
    return jsonify({"status": "ok"})


# ============================================================
# /extract — извлечение текста из DOCX (как раньше)
# ============================================================

@app.post('/extract')
def extract():
    try:
        if 'file' in request.files:
            data = request.files['file'].read()
        elif request.data:
            data = request.data
        else:
            return jsonify({"error": "No file provided"}), 400

        if not data:
            return jsonify({"error": "Empty file"}), 400

        doc = Document(io.BytesIO(data))

        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        text = '\n'.join(parts)
        log.info(f"Extracted {len(text)} chars from {len(data)} byte docx")

        return jsonify({
            "text": text,
            "chars": len(text),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        })

    except Exception as e:
        log.error(f"Extract error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


# ============================================================
# /generate-contract — рендер шаблона + конвертация в PDF
# ============================================================

def _convert_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    """Конвертирует .docx в PDF через LibreOffice headless. Возвращает путь к PDF."""
    result = subprocess.run(
        [
            'libreoffice', '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(out_dir),
            str(docx_path),
        ],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice failed: {result.stderr}")
    pdf_path = out_dir / (docx_path.stem + '.pdf')
    if not pdf_path.exists():
        raise RuntimeError(f"PDF не создан: {result.stdout} / {result.stderr}")
    return pdf_path


def _render_contract(payload):
    """
    Общая часть для PDF/DOCX-эндпоинтов:
    валидирует payload, рендерит docxtpl-шаблон,
    возвращает (tmp_dir_obj, docx_path, contract_number).
    """
    contract_number = payload.get('contract_number')
    contract_date   = payload.get('contract_date')
    client          = payload.get('client') or {}

    if not contract_number or not contract_date:
        raise ValueError("contract_number и contract_date обязательны")
    if not client.get('company_name'):
        raise ValueError("client.company_name обязателен")

    # Дефолты для опциональных полей
    for k in ('inn','kpp','ogrn','address','account','bank','bik','cor_account',
              'director_genitive','director_short','email','phone'):
        client.setdefault(k, '')

    context = {
        'contract_number': contract_number,
        'contract_date':   contract_date,
        'client':          client,
        'company':         COMPANY,
    }

    tmp_dir = tempfile.TemporaryDirectory()
    tmp = Path(tmp_dir.name)
    docx_path = tmp / f"contract_{contract_number}.docx"
    doc = DocxTemplate(TEMPLATE_PATH)
    doc.render(context)
    doc.save(docx_path)

    return tmp_dir, docx_path, contract_number


@app.post('/generate-contract')
def generate_contract():
    """POST JSON → PDF (application/pdf). Принимает {contract_number, contract_date, client:{...}}."""
    try:
        payload = request.get_json(force=True, silent=True) or {}
        tmp_dir, docx_path, contract_number = _render_contract(payload)

        with tmp_dir:
            pdf_path  = _convert_to_pdf(docx_path, Path(tmp_dir.name))
            pdf_bytes = pdf_path.read_bytes()
            log.info(f"Generated PDF for {contract_number} ({len(pdf_bytes)} bytes)")

            return send_file(
                io.BytesIO(pdf_bytes),
                mimetype='application/pdf',
                as_attachment=True,
                download_name=f"Contract_{contract_number}.pdf"
            )

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        log.error(f"Generate PDF error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@app.post('/generate-contract-docx')
def generate_contract_docx():
    """POST JSON → DOCX. Тот же контракт, но возвращается .docx без конвертации в PDF."""
    try:
        payload = request.get_json(force=True, silent=True) or {}
        tmp_dir, docx_path, contract_number = _render_contract(payload)

        with tmp_dir:
            docx_bytes = docx_path.read_bytes()
            log.info(f"Generated DOCX for {contract_number} ({len(docx_bytes)} bytes)")

            return send_file(
                io.BytesIO(docx_bytes),
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f"Contract_{contract_number}.docx"
            )

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        log.error(f"Generate DOCX error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8082, debug=False)
